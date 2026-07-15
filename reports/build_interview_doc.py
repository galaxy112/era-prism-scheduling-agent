from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("reports/interview_questions.docx")


SECTIONS = [
    (
        "一、论文理解类",
        [
            (
                "你怎么用一句话解释 ERA？",
                "ERA 是一个面向经验研究代码的 AI 辅助闭环：模型生成或改写代码，系统执行它，用结果和评分反馈继续搜索更好的实现。我的迁移不是复现树搜索，而是保留“候选 -> 执行 -> 评分 -> 改进”的核心机制。",
            ),
            (
                "你怎么用一句话解释 PRISM？",
                "PRISM 强调把实验协议结构化，并通过模拟验证协议是否可执行、是否有错误，再把反馈用于修复协议。迁移到排程里，就是把最佳排程转成结构化 step list，再由轻量模拟器复验资源、顺序和停机窗口。",
            ),
            (
                "为什么不完整复现论文？",
                "题目明确给了 1.5 天和 Demo 边界，评价重点是迁移能力和闭环验证。完整复现 ERA/PRISM 的工程量远超作业目标，所以我选择最小闭环，把关键思想讲清楚并跑通。",
            ),
            (
                "ERA 和普通启发式搜索有什么区别？",
                "普通启发式搜索通常在人写定的邻域里探索；ERA 的亮点是把候选代码本身也纳入生成和改写对象。我的 Demo 用离线模板模拟这一点，若继续扩展，可以让 LLM 真实生成候选函数并在沙箱中测试。",
            ),
        ],
    ),
    (
        "二、系统设计类",
        [
            (
                "你的系统闭环是什么？",
                "读取问题 JSON，运行候选算法，验证硬约束，计算分数，根据反馈生成或选择下一候选，保留最低分方案，最后转结构化协议并用模拟器复验。",
            ),
            (
                "为什么把验证器和评分器拆开？",
                "验证器负责事实判断：有没有冲突、顺序错误、停机冲突。评分器负责偏好权衡：硬错误、完工时间、延期时间分别多重要。拆开后更容易测试，也更容易替换目标函数。",
            ),
            (
                "为什么候选生成不用真实 LLM API？",
                "为了可复现、无密钥依赖、演示稳定。面试作业的核心是闭环机制而不是 API 集成。README 和 AI_USAGE 中也明确说明这是离线模板候选，后续可扩展为真实 LLM 沙箱生成。",
            ),
            (
                "PRISM 式模拟器为什么复用验证器？",
                "这个 Demo 的仿真目标是验证结构化协议的硬约束可执行性，因此复用同一套约束规则是合理的。更完整版本可以加入样本状态、设备准备时间、失败概率和实验结果模型。",
            ),
        ],
    ),
    (
        "三、算法与评分类",
        [
            (
                "基线算法是什么？",
                "基线按批次输入顺序贪心排程，每个步骤选择批次可用时间和资源可用时间的最大值作为开始时间。它故意不处理停机窗口，用来展示失败反馈。",
            ),
            (
                "候选算法有哪些？",
                "包括 naive_batch_order、repair_conflicts、downtime_aware_earliest 和 deadline_priority。它们分别代表初始候选、基于反馈修复、显式停机感知和搜索排序变化。",
            ),
            (
                "评分公式为什么这样设计？",
                "沿用题面建议：1000 * 硬约束错误数 + 总完工时间 + 3 * 总延期时间。硬错误权重很高，确保系统先追求可行，再优化效率和截止期。",
            ),
            (
                "你的结果是最优解吗？",
                "不保证。这个作业不要求数学最优，我的目标是得到可行且可解释的改进结果。若追求最优，可以建模为 job-shop scheduling，用 CP-SAT、MILP 或更系统的局部搜索求解。",
            ),
        ],
    ),
    (
        "四、验证与测试类",
        [
            (
                "验证器检查哪些硬约束？",
                "三类：同一资源同一时间只能服务一个步骤；同一批次步骤必须按顺序执行；步骤不能和设备停机窗口重叠。",
            ),
            (
                "你如何证明闭环真的发生？",
                "示例输出中第一个候选出现 downtime_conflict，第二个候选根据反馈修复后硬错误降为 0，第四个候选进一步降低分数。测试也断言后续候选优于失败候选。",
            ),
            (
                "如果模拟器发现协议失败怎么办？",
                "把模拟器返回的 issue 作为下一轮候选生成输入。当前 Demo 中协议复验通过；如果失败，可以像处理 validator feedback 一样触发 repair candidate。",
            ),
            (
                "测试覆盖了什么？",
                "覆盖资源冲突、步骤顺序错误、停机冲突、评分公式，以及闭环从失败候选改进到无硬错误候选的行为。",
            ),
        ],
    ),
    (
        "五、AI 工具使用与工程判断",
        [
            (
                "AI 在这个作业里帮了什么？",
                "AI 帮助拆解题目、形成架构、生成初版代码和文档。但算法边界、离线候选策略、测试验收和最终叙事需要人工判断。",
            ),
            (
                "你如何避免 AI 生成看似正确但不可运行的东西？",
                "把系统拆成小模块，写单元测试，运行 CLI 示例，把 README 的命令复制验证，并保留真实 sample_output。",
            ),
            (
                "如果 LLM 生成了错误候选，系统怎么处理？",
                "候选必须经过执行、验证和评分。错误候选不会直接进入结果，而是变成反馈信号，帮助下一轮修复或被更好候选淘汰。",
            ),
            (
                "你会如何把这个 Demo 变成生产系统？",
                "增加沙箱执行、候选代码安全检查、持久化实验记录、更强求解器、可视化排程、真实设备状态接口，以及对协议模拟器的领域建模。",
            ),
        ],
    ),
    (
        "六、继续扩展追问",
        [
            (
                "如果再给一周，你优先做什么？",
                "第一，接入真实 LLM 生成候选并沙箱测试；第二，增加 CP-SAT 或局部搜索；第三，把模拟器扩展成实验状态机；第四，做甘特图和错误定位 UI。",
            ),
            (
                "如果面试官质疑离线候选不算 ERA，你怎么回答？",
                "我会承认它不是完整 ERA，只是保留 ERA 的实验闭环骨架。选择离线候选是为了 Demo 稳定和可复现；接口已经可以替换为真实 LLM 生成候选。",
            ),
            (
                "如果面试官问 PRISM 体现得不够强，你怎么回答？",
                "当前 PRISM 体现为结构化协议和模拟验证。更强版本会模拟样本状态、设备准备/清洗、实验失败概率和协议语义检查，但这超出 1.5 天最小 Demo。",
            ),
            (
                "你最想强调的设计取舍是什么？",
                "优先保证可运行和可解释，而不是堆功能。这个题考的是判断边界、形成闭环、诚实说明限制，而不是在短时间内做一个庞大但不稳定的系统。",
            ),
        ],
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_question_table(doc: Document, questions: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False
    headers = table.rows[0].cells
    headers[0].text = "可能问题"
    headers[1].text = "回答要点"
    for cell in headers:
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for question, answer in questions:
        cells = table.add_row().cells
        cells[0].text = question
        cells[1].text = answer
        set_cell_width(cells[0], 3100)
        set_cell_width(cells[1], 6260)


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build() -> None:
    doc = Document()
    apply_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ERA + PRISM 实验室排程 Agent 面试追问准备")
    title_run.font.name = "Microsoft YaHei"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("用于二面答辩前快速复习：问题、回答要点与取舍说明")
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.color.rgb = RGBColor.from_string("555555")

    doc.add_heading("使用建议", level=1)
    doc.add_paragraph(
        "先把每个问题的第一句话背熟，再根据面试官追问展开细节。回答时要主动承认 Demo 边界：它展示的是论文思想迁移和闭环验证，不是完整论文复现。"
    )

    for heading, questions in SECTIONS:
        doc.add_heading(heading, level=1)
        add_question_table(doc, questions)

    doc.add_heading("一句话总结", level=1)
    doc.add_paragraph(
        "这个作业我优先保证可运行、可解释、可验证：用 ERA 的候选改进闭环解决排程搜索，用 PRISM 的结构化协议和模拟验证确保结果能被机器检查。"
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
